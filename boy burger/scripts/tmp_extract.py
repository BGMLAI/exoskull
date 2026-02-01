import os, re, json
from pathlib import Path
from docx import Document
from pypdf import PdfReader
from PIL import Image
import pytesseract

ROOT = Path('.')
EXTS = {'.md', '.docx', '.pdf', '.url', '.png', '.jpg', '.jpeg', '.gif', '.webp'}
MAX_OCR_PIXELS = 8000000


def read_text(path: Path):
    ext = path.suffix.lower()
    if ext == '.md':
        return path.read_text(encoding='utf-8', errors='ignore')
    if ext == '.url':
        return path.read_text(encoding='utf-8', errors='ignore')
    if ext == '.docx':
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return '\n'.join(parts)
    if ext == '.pdf':
        try:
            reader = PdfReader(str(path))
            texts = []
            for page in reader.pages:
                t = page.extract_text() or ''
                if t:
                    texts.append(t)
            return '\n'.join(texts)
        except Exception as e:
            return f"[PDF_READ_ERROR] {e}"
    if ext in {'.png', '.jpg', '.jpeg', '.gif', '.webp'}:
        try:
            with Image.open(path) as img:
                w, h = img.size
                if w * h > MAX_OCR_PIXELS:
                    img = img.resize((int(w*0.5), int(h*0.5)))
                return pytesseract.image_to_string(img, lang='eng+pol')
        except Exception as e:
            return f"[OCR_ERROR] {e}"
    return ''

files = []
for p in ROOT.rglob('*'):
    if not p.is_file():
        continue
    if p.suffix.lower() in EXTS:
        if any(part.startswith('.') for part in p.parts):
            continue
        files.append(p)

out_path = ROOT / 'tmp_extracted_text.jsonl'
summary_path = ROOT / 'tmp_extracted_summary.md'

fact_patterns = [
    r"\b\d{4}\b",
    r"\b\d+[\.,]?\d*%\b",
    r"\bN\s*=\s*\d+\b",
    r"\bp\s*[<=>]\s*0?\.\d+\b",
    r"meta-analiz|meta analysis|systematic review|randomized|RCT|longitudinal|cohort|case-control|przegląd",
    r"\b(UNESCO|OECD|GUS|WHO|EUROSTAT|CDC|ONS|APA|APA 7)\b",
]
import re
fact_re = re.compile('|'.join(fact_patterns), re.IGNORECASE)

with out_path.open('w', encoding='utf-8') as outf, summary_path.open('w', encoding='utf-8') as sumf:
    sumf.write('# TMP EXTRACTED FACT CANDIDATES\n\n')
    for p in sorted(files):
        rel = p.relative_to(ROOT)
        try:
            text = read_text(p)
        except Exception as e:
            text = f"[READ_ERROR] {e}"
        outf.write(json.dumps({'path': str(rel), 'text': text[:200000]}, ensure_ascii=False) + '\n')
        candidates = []
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            if fact_re.search(line):
                candidates.append(line)
            if len(candidates) >= 12:
                break
        if candidates:
            sumf.write(f"## {rel}\n")
            for c in candidates:
                c = c.replace('\t', ' ')
                if len(c) > 300:
                    c = c[:300] + '…'
                sumf.write(f"- {c}\n")
            sumf.write('\n')

print(f"Processed {len(files)} files. Output: {out_path}, {summary_path}")
